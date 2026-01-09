import json, os, pdfplumber, re, base64, io, openpyxl
import pandas as pd
from pdf2docx import parse
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from openpyxl.drawing.image import Image as OpenPyXlImage
PRINT_LOG = False
CLEAN_AFTER_EXECUTION = False
def find_element(children, target):
    for child in children:
        if child["type"] == "heading":
            if child["text"] == target:
                return child["children"]
            if find_element(child["children"], target):
                return find_element(child["children"], target)
    return None

def combine_para_children(child_list):
    combined = ""
    for child in child_list:
        if child["type"] == "paragraph":
            combined += child["text"]
    return combined

def find_all_imgs(children):
    imgs = []
    for child in children:
        if child["type"] == "image":
                imgs.append(child)
        elif child["type"] == "heading":
            imgs += find_all_imgs(child["children"])
    return imgs

def find_all_tables(children):
    tabs = []
    for child in children:
        if child["type"] == "table":
            tabs.append(child)
        elif child["type"] == "heading":
            tabs += find_all_tables(child["children"])
    return tabs

def merge_by_title(items):
    merged = {}
    order = []

    for item in items:
        title = item["title"]
        label = item["label"]
        content = item.get("content", [])

        if title not in merged:
            merged[title] = {
                "title": title,
                "label": label,
                "content": []
            }
            order.append(title)

        merged[title]["content"].extend(content)

    return [merged[t] for t in order]


def geo_layers_spec():
    wb = openpyxl.load_workbook("./output/plumber_temp.xlsx")
    ws = None
    tab_holder = []
    for sheet in wb:
        if sheet["A1"].value and "地层名称" in sheet["A1"].value:
            ws = sheet
    if ws:
        for row_num in range(1, ws.max_row + 1):
            row_values = [cell.value for cell in ws[row_num]]
            tab_holder.append(row_values)
    else:
        return None
    def is_number(s):
        try:
            float(s)
            return True
        except ValueError:
            return False
    for i, row in enumerate(tab_holder):
        for j, cell in enumerate(row):
            if cell:
                splited = tab_holder[i][j].replace(".\n", ".").split("\n")
                if len(splited) > 1 and is_number(splited[0]) and is_number(splited[1]):
                    tab_holder[i][j] = splited[0]
                    tab_holder[i + 1][j] = splited[1]
    for i, row in enumerate(tab_holder):
        for j, cell in enumerate(row):
            if cell:
                tab_holder[i][j] = tab_holder[i][j].replace("\n","").replace(" ","")
    return tab_holder

def excel_generation(data_dict):
    wb = openpyxl.Workbook()
    
    doc_dict = data_dict['children']
    
    tablist = merge_by_title(find_all_tables(doc_dict))
    img_list = find_all_imgs(doc_dict)
    
    # 表格提取
    for tab in tablist:
        ws = wb.create_sheet(tab['title'])
        if "地层分层" in tab['title']:
            geo_layers_tab = geo_layers_spec()
            if geo_layers_tab:
                for row in geo_layers_tab:
                    ws.append(row)
                continue
        for row in tab['content']:
            ws.append(row)
    
    # 图片提取
    for img in img_list:
        img_data = base64.b64decode(img['data'])
        ws = wb.create_sheet(f"{img['title']}")
        img_stream = io.BytesIO(img_data)
        excel_img = OpenPyXlImage(img_stream)
        ws.add_image(excel_img, "A1")
    
    # 内容提取
    for heading in doc_dict:
        if not heading["type"] == "heading":
            continue
        try:
            ws = wb.create_sheet(f"{heading['number']} {heading['text']}")
        except ValueError:
            print(f"{heading['number']} {heading['text']}")
            continue
            
        def construct_list(node):
            if node["type"] == "table" or node["type"] == "image":
                return
            if node['type'] == "heading":
                ws.append([f"{node['number']} {node['text']}"])
                for child in node['children']:
                    construct_list(child)
            else:
                ws.append([node["text"]])
        construct_list(heading)
    
    wb.remove(wb["Sheet"])
    # Save the workbook
    wb.save(f"./output/extracted_data_base.xlsx")



def search():
    with open('./output/output.json', 'r', encoding="utf-8-sig") as file:
        data_dict = json.load(file)

    excel_generation(data_dict)
    
    with open(f"./output/extracted_data_base.xlsx", 'rb') as file:
            binary_data = file.read()
            base64_encoded_data = base64.b64encode(binary_data)
            return base64_encoded_data.decode('utf-8')

def replace_linebreaks_with_paragraphs_simple(input_path, output_path):
    doc = Document(input_path)
    
    for para in list(doc.paragraphs):
        parts = para.text.split("\n")
        if len(parts) > 1:
            parent = para._element.getparent()
            index = parent.index(para._element)  # save index before removing
            parent.remove(para._element)
            
            for i, part in enumerate(parts):
                new_para = doc.add_paragraph(part, style=para.style)
                # insert new paragraph at the correct position
                parent.insert(index + i, new_para._element)
    
    doc.save(output_path)

def iter_block_items(parent):
    """
    Yield paragraphs and tables in document order.
    (Images are inside paragraphs, so we detect them separately.)
    """
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if child.tag.endswith('p'):
            yield Paragraph(child, parent)
        elif child.tag.endswith('tbl'):
            yield Table(child, parent)

def paragraph_to_base64_image(paragraph):
    """
    If the paragraph contains an image, return its Base64 string.
    Otherwise, return None.
    """
    for run in paragraph.runs:
        blips = run.element.xpath('.//a:blip')
        for blip in blips:
            rId = blip.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            image_part = paragraph.part.related_parts[rId]
            image_bytes = image_part.blob
            return base64.b64encode(image_bytes).decode("utf-8")
    return None

def classify_doc_blocks(docx_path):
    """
    Iterate through a DOCX and classify blocks as paragraph, table, or image.
    For images, return the Base64 string.
    """
    waiting_image_label = False
    waiting_table_content = False
    image_data_holder = ""
    table_label_holder = ""
    table_text_holder = ""
    doc = Document(docx_path)
    root = {"type": "document", "children": []}
    stack = [root]  # hierarchy stack
    heading_pattern = re.compile(r"^(\d+[.\u3001]*\d*(?:[.\u3001]\d*)*)\s*(.*)$")
    table_pattern = re.compile(r"^(表)\s*(\d+-\d+)\s*(.*)$")
    image_pattern = re.compile(r"^(图)\s*(\d+-\d+)\s*(.*)$")
    number_holder = "-1"
    
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            img_b64 = paragraph_to_base64_image(block)
            if img_b64: # image
                image_data_holder = img_b64
                waiting_image_label = True
            else: # text paragraph
                text = block.text.strip()
                if not text:
                    continue
                match = heading_pattern.match(text)
                if match:
                    if re.compile(r"^\d+$").match(text):
                        # Pure number paragraph (probably page number)
                        continue
                    # heading paragraph
                    numbering = match.group(1)   # e.g. "1.2.3"
                    if (numbering[-1] == "、"):
                        numbering = numbering[:-1]
                    title_text = match.group(2)  # e.g. "Background"
                    if valid_numbering(number_holder, numbering):
                        number_holder = numbering
                        level = numbering.count(".") + 1

                        node = {
                            "type": "heading",
                            "level": level,
                            "number": numbering,
                            "text": title_text,
                            "children": []
                        }

                        # Fix stack hierarchy
                        while len(stack) > 1 and stack[-1].get("level", 0) >= level:
                            stack.pop()

                        stack[-1]["children"].append(node)
                        stack.append(node)
                        waiting_table_content = False
                else:
                    if re.compile(r"^\d+$").match(text):
                        # Pure number paragraph (probably page number)
                        continue
                    tab_match = table_pattern.match(text)
                    img_match = image_pattern.match(text)
                    if tab_match:
                        # Table label
                        table_label_holder = tab_match.group(2)
                        table_text_holder = tab_match.group(3)
                        waiting_table_content = True
                    elif img_match and waiting_image_label:
                        # Image label
                        node = {
                            "type": "image",
                            "label": img_match.group(2),
                            "title": img_match.group(3),
                            "data": image_data_holder,
                        }
                        stack[-1]["children"].append(node)
                        waiting_image_label = False
                    else:
                        # Normal paragraph
                        node = {"type": "paragraph", "text": text}
                        stack[-1]["children"].append(node)
                        waiting_table_content = False
                    
        elif isinstance(block, Table):
            if not waiting_table_content:
                continue
            # Table
            tab = extract_table_content(block)
 
            if tab[0][0] == "中国石化" and (tab[0][2] == "钻井地质设计" or tab[0][2] == "钻井工程设计"):
                continue
            def convert_if_number(s):
                try:
                    return int(s)
                except ValueError:
                    try:
                        return float(s)
                    except ValueError:
                        return s
            for i, a in enumerate(tab):
                for j, b in enumerate(a):
                    tab[i][j] = convert_if_number(b)
            node = {"type": "table", "label": table_label_holder, "title": table_text_holder, "content": tab}
            stack[-1]["children"].append(node)

    return root

def extract_table_content(table):
    """
    Extract all text from a python-docx Table as a 2D list of strings.
    Each cell's text is joined if it has multiple paragraphs.
    """
    table_data = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            # A cell may have multiple paragraphs, join them with newlines
            cell_text = "\n".join([p.text for p in cell.paragraphs])
            row_data.append(cell_text)
        table_data.append(row_data)
    return table_data

def valid_numbering(last_heading, current_heading):
    """
    Determine if current_heading is valid given last_heading.

    Examples:
        valid_numbering("1.2", "1.3") -> True
        valid_numbering("1.2", "2.1") -> True
        valid_numbering("1.2", "1.4") -> False
        valid_numbering("1.2", "3.1") -> False
    """
    if (last_heading == "-1"):
        return True
    if current_heading[0] == "0":
        return False
    # Split numbering into integer lists
    try:
        last_nums = [int(x) for x in last_heading.split(".")]
        curr_nums = [int(x) for x in current_heading.split(".")]
    except ValueError:
        return False  # non-integer parts are invalid

    # If current heading has fewer levels, only compare common prefix
    min_len = min(len(last_nums), len(curr_nums))

    # Check if prefix matches except last element
    if last_nums[:min_len-1] != curr_nums[:min_len-1]:
        return False

    # Now check the last number
    last_num = last_nums[min_len-1]
    curr_num = curr_nums[min_len-1]

    # Valid if:
    # 1. Same level and increment by 1
    # 2. New sublevel (current level = last level +1) starts with 1
    if len(curr_nums) == len(last_nums):
        return curr_num == last_num + 1
    elif len(curr_nums) == len(last_nums) + 1:
        return True
    elif len(curr_nums) < len(last_nums):
        # Going back to upper level
        return curr_num == last_num + 1
    else:
        # Jumping more than one level deeper → invalid
        return False
    
def extraction(string_input, input_type="file_name"):
    
    if (input_type=="file_name"):
        parse(string_input,"./output/input.docx")
    elif (input_type == "base64"):
        # Decode the Base64 PDF string
        try:
            pdf_binary = base64.b64decode(string_input)
        except:
            return "Invalid Base64 string"
        
        # Create a temporary PDF file to hold the decoded data
        with open("./output/temp.pdf", "wb") as f:
            f.write(pdf_binary)
            
        parse("./output/temp.pdf","./output/input.docx")

    replace_linebreaks_with_paragraphs_simple("./output/input.docx", "./output/output.docx")

def splitPDFpages(target):
    from PyPDF2 import PdfWriter, PdfReader
    print(f"Removing Pages From: {target}")
    def find_visual_toc_pages(filepath):
        toc_pages = [0]
        with open(filepath, 'rb') as file:
            reader = PdfReader(file)
            num_pages = len(reader.pages)

            # Start checking from the second page (index 1)
            for i in range(1, num_pages):
                page = reader.pages[i]
                text = page.extract_text()
                if text:
                    # Further heuristic: Check for lines that end with a number (common in TOCs)
                    lines = text.strip().split('\n')
                    # potential_toc_lines = [line for line in lines if line.strip() and line.strip()[-1].isdigit()]
                    potential_toc_lines = []
                    for line in lines:
                        if len(line.strip()) == 0:
                            continue
                        if line.strip().count(".") / len(line.strip()) > 0.5:
                            potential_toc_lines.append(line)
                    # If a significant portion of lines look like TOC entries
                    if len(potential_toc_lines) / len(lines) > 0.3: 
                        # Using 1-based indexing for the user
                        toc_pages.append(i)
                        print(f"Found potential TOC on page {i + 1} (based on keywords and structure).")
                        continue # Move to next page as TOCs can span multiple pages
                    else:
                        break
        return toc_pages
    
    pages_to_delete = find_visual_toc_pages(target)
    infile = PdfReader(target)
    output_content = PdfWriter()
    output_title = PdfWriter()

    for i in range(len(infile.pages)):
        if i not in pages_to_delete:
            p = infile.pages[i]
            output_content.add_page(p)
        else:
            p = infile.pages[i]
            output_title.add_page(p)

    with open('./output/content.pdf', 'wb') as f:
        output_content.write(f)
    with open('./output/title.pdf', 'wb') as f:
        output_content.write(f)

def pdf_to_excel_pdfplumber(pdf_path):
    excel_path = "./output/plumber_temp.xlsx"
    try:
        with pdfplumber.open(pdf_path) as pdf:
            with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
                # 提取表格
                table_count = 0
                for page_num, page in enumerate(pdf.pages, 1):
                    tables = page.extract_tables()

                    for i, table in enumerate(tables):
                        if table:
                            # 清理表格数据
                            cleaned_table = []
                            for row in table:
                                cleaned_row = [str(cell).strip() if cell is not None else "" for cell in row]
                                # 跳过空行
                                if any(cleaned_row):
                                    cleaned_table.append(cleaned_row)

                            if cleaned_table:
                                # 创建DataFrame
                                df = pd.DataFrame(cleaned_table[1:], columns=cleaned_table[0])
                                sheet_name = f"Page{page_num}_Table{i + 1}"
                                df.to_excel(writer, sheet_name=sheet_name, index=False)
                                table_count += 1
        return True

    except Exception as e:
        print(f"Excel转换过程中出现错误: {e}")
        return False

def delete_folder(folder_path):
        import shutil
        # 检查文件夹是否存在
        if os.path.exists(folder_path):
            try:
                shutil.rmtree(folder_path)
            except Exception as e:
                # 捕获权限不足、文件被占用等异常
                print(f"删除失败：{e}")
        else:
            print(f"文件夹 {folder_path} 不存在")

def internal_calling(initial_PDF = None):
    if not initial_PDF:
        initial_PDF = "./output/design_file.pdf"
    os.makedirs("./output", exist_ok=True)
    # 切分
    splitPDFpages(initial_PDF)
    extraction("./output/content.pdf")
    extraction("./output/title.pdf")
    
    # 提取为json
    json_obj = classify_doc_blocks("./output/output.docx")
    with open('./output/output.json', 'w', encoding="utf-8") as f:
        f.write(json.dumps(json_obj, ensure_ascii=False, indent=2))

    # 复杂结构表格提取
    pdf_to_excel_pdfplumber(initial_PDF)
    
    # 生成excel
    excel_b64= search()
    
    if CLEAN_AFTER_EXECUTION:
        delete_folder("./output")
    
    return excel_b64

def extraction_enter_point(input_b64):
    """
    Args:
        input_b64: base64表示的PDF文件
    Returns:
        excelb64: excel文件的base64编码
    """

    os.makedirs("./output", exist_ok=True)
    with open('./output/design_file.pdf', 'wb') as f:
        f.write(base64.b64decode(input_b64))
    return internal_calling()

def extraction_entry_stream(input_binary):
    """
    Args:
        input_binary: 二进制表示的PDF文件
    Returns:
        excelb64: excel文件的base64编码
    """
    os.makedirs("./output", exist_ok=True)

    with open('./output/design_file.pdf', 'wb') as f:
        f.write(input_binary)
    return internal_calling()


# 示例调用
if __name__ == "__main__":
    # print(geo_layers_spec())
    excelb64 = internal_calling("./CDC913-X89/CDC913-X89钻井地质设计.pdf")